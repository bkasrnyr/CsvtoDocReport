import os
import sys

# Check if the required modules are installed
try:
    import pandas
    import openpyxl
    import tqdm
    import colorama
except ImportError:
    # Modules are not installed, install them using requirements.txt
    print("Modules are not installed, installing required modules using requirements.txt...")
    requirements_file = "requirements.txt"

    try:
        # Install the required modules using pip
        os.system(f"python -m pip install -r {requirements_file}")
    except Exception as e:
        print(f"Error installing modules: {e}")
        sys.exit(1)

    # Verify installation
    try:
        import pandas
        import openpyxl
        import tqdm
        import colorama
    except ImportError as e:
        print(f"Error importing modules: {e}")
        sys.exit(1)
try:
    import pandas as pd
    import numpy as np
    from docx import Document
    from docx.shared import Inches
    from docx.shared import RGBColor
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    import openpyxl
    import glob
    from pyexcel.cookbook import merge_all_to_a_book
    import warnings
    from colorama import Fore, Style
    from tqdm import tqdm
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError as e:
        print(Fore.RED + f"\n[-] {e}, please install using \033[0;35m python -m pip install -r requirements.txt\033[0m" + Style.RESET_ALL)
        exit()

def print_logo():
    logo = """
    ███████           ██                  ████████                                    ██  
  ██     ██                               ██    ██          ██████                    ██  
 ██       ██  ██   ██ ██  █████  ██  ██   ██    ██   █████  ██   ██  ██████  ██████ ██████
 ██       ██  ██   ██ ██ ██   ██ ██ ██    ███████   ██   ██ ██   ██ ██    ██  ██  █   ██ 
 ██    ██ ██  ██   ██ ██ ██      ████     ██   ██   ███████ ██████  ██    ██  ██      ██  
  ██     ██   ██   ██ ██ ██   ██ ██ ██    ██    ██  ██      ██      ██    ██  ██      ██  
   ███████ ██  ██████ ██  █████  ██  ██   ██     ██  ██████ ██       ██████  ███      █████

                                    \033[34m[✔] https://github.com/bkasrnyr/CsvtoDocReport  [✔]
                                    \033[34m[✔]            Version 1.0                      [✔]
                                    \033[34m[✔]       Author:Bikash Rouniyar                [✔]
                                    \033[91m[X] Please Don't Use For illegal Activity       [X]
 """
    print(logo + '\n')    

def merge_csv_to_xlsx(filename):
    merge_all_to_a_book(glob.glob(filename), "output.xlsx")
    wb1 = openpyxl.load_workbook('output.xlsx')
    ws1 = wb1.active
    ws1.delete_rows(1, 7)
    new_xlsx_name = "Output_XLSX_" + filename.strip(".csv") + ".xlsx"
    wb1.save(new_xlsx_name)
    os.remove("output.xlsx")
    return new_xlsx_name

def process_data(filename):
    df = pd.read_excel(filename, index_col=None) 
    df = df.assign(column=np.nan)
    df[['CVSS', 'Base']] = df['CVSS Base'].str.split(expand=True)
    df1 = df['CVSS'].fillna('-1.0', inplace=True)
    df2 = df.fillna('-', inplace=True)
    df[['IP1', 'IP2', 'IP3', 'IP4']] = df['IP'].str.split(".", n=3, expand=True).astype('int64')
    dfnn = df.sort_values(['IP1', 'IP2', 'IP3', 'IP4', 'CVSS'], ascending=[True, True, True, True, False])
    dfl = dfnn[df['Severity'] != 'None']
    nor = len(dfl)
    index_list = []
    for x in range(nor):
        index_list.append(x)
    dfl.insert(0, 'index', index_list)
    global dfn
    dfn = dfl.set_index('index')  
    return dfn

def create_docx_file(filename):
    empty_docx = Document("./DUMP/Converted.docx")
    global new_docx_name
    new_docx_name = "Converted" + filename.strip(".csv") + ".docx"
    empty_docx.save(new_docx_name)
    return new_docx_name

def write_csv_data_to_docx(df, docx_filename):
    nl = len(dfn['Title'])
    lol = [[] for _ in range(nl)]
    n2 = len(lol)
    z = 0
    for x in dfn['Title']:
        lol[z].append(z)
        z = z + 1
    uef = []
    document = Document(new_docx_name)
    m = 0
    p = 0
    progress_bar = tqdm(total=len(lol),bar_format="{l_bar}{bar}| \033[0;35m{n_fmt}/{total_fmt}\033[0m",ncols=80)
    try:
        for x in lol:
            progress_bar.set_description("\033[92m[+] Progress:\033[0m")
            p += 1
            for y in x:
                uef.append(y)
                break
        # *************************************************************************
        def make_rows_bold(*rows):
            for row in rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

        def remove_row(table, row):
            tbl = summarytable._tbl
            tr = row._tr
            tbl.remove(tr)
        # print '******************************************************************'

        summarytable = document.add_table(rows=2 + len(lol), cols=7)
        summarytable.style = 'TableGrid'
        column_widths = [0.3, 1.248031, 1.740157, 0.6929134, 0.5590551, 0.9409449, 1.748031]

        for col, width in enumerate(column_widths):
            summarytable.cell(0, col).width = Inches(width)
            summarytable.cell(1, col).width = Inches(width)

        header_texts = ['S. No.', 'Component', 'Vulnerabilities Noted per Component',
                        'Severity Level', 'CVSS Level', 'Compliance Status',
                        'Exceptions, False Positives, or Compensating Controls (Noted by the ASV for this Vulnerability)']

        for col, header_text in enumerate(header_texts):
            summarytable.cell(1, col).text = header_text

        shading_elm = parse_xml(r'<w:shd {} w:fill="#D9D9D9"/>'.format(nsdecls('w')))
        summarytable.cell(1,0)._tc.get_or_add_tcPr().append(shading_elm)
        shading_elm = parse_xml(r'<w:shd {} w:fill="#D9D9D9"/>'.format(nsdecls('w')))
        summarytable.cell(1,1)._tc.get_or_add_tcPr().append(shading_elm)
        shading_elm = parse_xml(r'<w:shd {} w:fill="#D9D9D9"/>'.format(nsdecls('w')))
        summarytable.cell(1,2)._tc.get_or_add_tcPr().append(shading_elm)
        shading_elm = parse_xml(r'<w:shd {} w:fill="#D9D9D9"/>'.format(nsdecls('w')))
        summarytable.cell(1,3)._tc.get_or_add_tcPr().append(shading_elm)
        shading_elm = parse_xml(r'<w:shd {} w:fill="#D9D9D9"/>'.format(nsdecls('w')))
        summarytable.cell(1,4)._tc.get_or_add_tcPr().append(shading_elm)
        shading_elm = parse_xml(r'<w:shd {} w:fill="#D9D9D9"/>'.format(nsdecls('w')))
        summarytable.cell(1,5)._tc.get_or_add_tcPr().append(shading_elm)
        shading_elm = parse_xml(r'<w:shd {} w:fill="#D9D9D9"/>'.format(nsdecls('w')))
        summarytable.cell(1,6)._tc.get_or_add_tcPr().append(shading_elm)
        a = summarytable.cell(0, 0)
        b = summarytable.cell(0, 1)
        c = summarytable.cell(0, 2)
        d = summarytable.cell(0, 3)
        e = summarytable.cell(0, 4)
        f = summarytable.cell(0, 5)
        g = summarytable.cell(0, 6)
        G = a.merge(b)
        H = c.merge(d)
        I = e.merge(f)
        J = G.merge(H)
        K = I.merge(J)
        L = K.merge(g)
        shading_elm2 = parse_xml(r'<w:shd {} w:fill="#595959"/>'.format(nsdecls('w')))
        L.text = ''
        run = L.paragraphs[0].add_run('Part 3a. Vulnerabilities Noted for each Component')
        font = run.font
        # c = Color("blue")
        font.color.rgb = RGBColor(255, 255, 255)
        L._tc.get_or_add_tcPr().append(shading_elm2)
        make_rows_bold(summarytable.rows[1])
        document.add_page_break()

        table = ['x'] * (uef[-1] + 1)

        # print uef
        string_host = [[]] * len(uef)

        strg = ''
        p = 0
        for y in lol:

            for z in y:
                string_host[p] = string_host[p] + [(str(
                    dfn['IP'][z] + ' ' + str('Port') + ' ' + str(dfn['Port'][z]) + '' + str('/') + '' + str(
                        dfn['Protocol'][z])))]
            p = p + 1
        q = 0
        g = 0
        counting = 2
        strg = [[]]
        for x in string_host:
            strg.append([list(set(x))])
        t = 1
        pp = 0
        for x in uef:

            c = 0
            strgf = ''.join(str(e) for e in strg[t])
            # ttrgf=''.join(str(e) for e in ttrg[t])

            # if df1['Risk'][x] != "None":
            table[x] = document.add_table(rows=16, cols=2)
            table[x].style = 'TableGrid'
            table[x].autofit = False
            column_width = Inches(1.57)

            for i in range(16):
                table[x].cell(i, 0).width = column_width

            table[x].cell(0, 0).text = str(pp + 1)
            table[x].cell(0, 1).text = str(dfn['Title'][x])
            table[x].cell(1, 0).text = 'PCI Compliance'
            table[x].cell(2, 0).text = 'PCI Severity'
            table[x].cell(3, 0).text = 'CVSS Base Score'
            table[x].cell(4, 0).text = 'Affected Instances'
            table[x].cell(5, 0).text = 'FQDN/Host Name'
            table[x].cell(6, 0).text = 'Vulnerability Type'
            table[x].cell(7, 0).text = 'Operating System'
            table[x].cell(8, 0).text = 'Category'
            table[x].cell(9, 0).text = 'CVE ID'
            table[x].cell(10, 0).text = 'Vendor Reference'
            table[x].cell(11, 0).text = 'Threat'
            table[x].cell(12, 0).text = 'Impact'
            table[x].cell(13, 0).text = 'Solution'
            table[x].cell(14, 0).text = 'Result'
            table[x].cell(15, 0).text = ''
            cell_width = Inches(5.751969)

            for i in range(16):
                table[x].cell(i, 1).width = cell_width

            row = table[x].rows
            row.height = Inches(0.1220472)
            A = table[x].cell(15, 0).merge(table[x].cell(15, 1))
            shading_elm = parse_xml(r'<w:shd {} w:fill="#D9D9D9"/>'.format(nsdecls('w')))
            A._element.get_or_add_tcPr().append(shading_elm)
            make_rows_bold(table[x].rows[0], table[x].rows[1], table[x].rows[2], table[x].rows[3], table[x].rows[4],
                           table[x].rows[5], table[x].rows[6], table[x].rows[7], table[x].rows[8], table[x].rows[9],
                           table[x].rows[10], table[x].rows[11], table[x].rows[12], table[x].rows[13], table[x].rows[14])
            document.add_page_break()
            progress_bar.update(1)
    except KeyboardInterrupt:
        progress_bar.close()
        print(Fore.RED + "\n[-] Program interrupted by the user. Exiting..." + Style.RESET_ALL)
        sys.exit(1)    
        if str(float(dfn['CVSS'][x])) != str('-1.0'):
            summarytable.cell(m + 2, 0).text = str(m + 1)
            summarytable.cell(m + 2, 1).text = strgf.replace("[", "").replace("]", "").replace("'", "").replace("Port -/-", "").replace(".0/tcp", "/tcp").replace(".0/udp", "/udp")
            summarytable.cell(m + 2, 2).text = str(dfn['Title'][x])
            v = np.array([0, 1, 1.1, 1.2, 1.3, 1.4, 1.5, 1.6, 1.7, 1.8, 1.9, 2, 2.1, 2.2, 2.3, 2.4, 2.5, 2.6, 2.7, 2.8, 2.9, 3, 3.1,
                  3.2, 3.3, 3.4, 3.5, 3.6, 3.7, 3.8, 3.9]).astype(str)

            w = np.array([4, 4.1, 4.2, 4.3, 4.4, 4.5, 4.6, 4.7, 4.8, 4.9, 5, 5.1, 5.2, 5.3, 5.4, 5.5, 5.6, 5.7, 5.8, 5.9, 6, 6.1,
                          6.2, 6.3, 6.4, 6.5, 6.6, 6.7, 6.8, 6.9]).astype(str)

            u = np.array([7, 7.1, 7.2, 7.3, 7.4, 7.5, 7.6, 7.7, 7.8, 7.9, 8, 8.1, 8.2, 8.3, 8.4, 8.5, 8.6, 8.7, 8.8, 8.9, 9, 9.1,
                          9.2, 9.3, 9.4, 9.5, 9.6, 9.7, 9.8, 9.9, 10]).astype(str)

            cvss_images = {
                str(float(dfn['CVSS'][x])): './Images/High.png' if str(float(dfn['CVSS'][x])) in u
                else './Images/Medium.png' if str(float(dfn['CVSS'][x])) in w
                else './Images/Low.png' if str(float(dfn['CVSS'][x])) in v
                else './Images/Informational.png'
            }
            summarytable.cell(m + 2, 3).text = ''
            p = summarytable.cell(m + 2, 3).paragraphs[0]
            r = p.add_run()
            r.add_picture(cvss_images.get(str(float(dfn['CVSS'][x])), './Images/Informational.png'))
            counting = counting + 1
            if float(dfn['CVSS'][x]) == -1.0:
                summarytable.cell(m + 2, 4).text = '-'
            else:
                CC = str(dfn['CVSS'][x]).replace(".0", "")
                summarytable.cell(m + 2, 4).text = CC

            if str(dfn['PCI Vuln'][x]) == str('yes'):
                summarytable.cell(m + 2, 5).text = str("Fail")
                shading_elm1 = parse_xml(r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls('w')))
                summarytable.cell(m + 2, 5)._tc.get_or_add_tcPr().append(shading_elm1)
            else:
                summarytable.cell(m + 2, 5).text = str("Pass")
                shading_elm1 = parse_xml(r'<w:shd {} w:fill="#92D050"/>'.format(nsdecls('w')))
                summarytable.cell(m + 2, 5)._tc.get_or_add_tcPr().append(shading_elm1)
            Column = str(dfn['column'][x]).replace("NaN", "")
            summarytable.cell(m + 2, 6).text = Column


            m = m + 1
        else:
            row = summarytable.rows[counting]
            remove_row(table, row)

        for col in table[x].columns:
            if c == 0:
                c = +1
                continue
            if str(dfn['PCI Vuln'][x]) == str('yes'):
                col.cells[1].text = str("Fail")
                shading_elm = parse_xml(r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls('w')))
                col.cells[1]._tc.get_or_add_tcPr().append(shading_elm)
            else:
                col.cells[1].text = str("Pass")
                shading_elm = parse_xml(r'<w:shd {} w:fill="#92D050"/>'.format(nsdecls('w')))
                col.cells[1]._tc.get_or_add_tcPr().append(shading_elm)
            v = np.array([0, 1, 1.1, 1.2, 1.3, 1.4, 1.5, 1.6, 1.7, 1.8, 1.9, 2, 2.1, 2.2, 2.3, 2.4, 2.5, 2.6, 2.7, 2.8, 2.9, 3, 3.1,
                  3.2, 3.3, 3.4, 3.5, 3.6, 3.7, 3.8, 3.9]).astype(str)

            w = np.array([4, 4.1, 4.2, 4.3, 4.4, 4.5, 4.6, 4.7, 4.8, 4.9, 5, 5.1, 5.2, 5.3, 5.4, 5.5, 5.6, 5.7, 5.8, 5.9, 6, 6.1,
                          6.2, 6.3, 6.4, 6.5, 6.6, 6.7, 6.8, 6.9]).astype(str)

            u = np.array([7, 7.1, 7.2, 7.3, 7.4, 7.5, 7.6, 7.7, 7.8, 7.9, 8, 8.1, 8.2, 8.3, 8.4, 8.5, 8.6, 8.7, 8.8, 8.9, 9, 9.1,
                          9.2, 9.3, 9.4, 9.5, 9.6, 9.7, 9.8, 9.9, 10]).astype(str)

            cvss_images = {
                str(float(dfn['CVSS'][x])): './Images/High.png' if str(float(dfn['CVSS'][x])) in u
                else './Images/Medium.png' if str(float(dfn['CVSS'][x])) in w
                else './Images/Low.png' if str(float(dfn['CVSS'][x])) in v
                else './Images/Informational.png'
            }
            col.cells[2].text = ''
            p = col.cells[2].paragraphs[0]
            r = p.add_run()
            r.add_picture(cvss_images.get(str(float(dfn['CVSS'][x])), './Images/Informational.png'))
            col.cells[3].text = '-' if float(dfn['CVSS'][x]) == -1.0 else str(dfn['CVSS'][x]).replace(".0", "") + ', ' + str(dfn['Base'][x]).replace("-","")
            col.cells[4].text = strgf.replace("[", "").replace("]", "").replace("'", "").replace("Port -/-", "").replace(".0/tcp", "/tcp").replace(".0/udp", "/udp")
            col.cells[5].text = str(dfn['DNS'][x])
            type_mapping = {'Practice': 'Potential', 'Vuln': 'Confirmed'}
            col.cells[6].text = type_mapping.get(str(dfn['Type'][x]), 'Informational')
            col.cells[7].text = str(dfn['OS'][x])
            col.cells[8].text = str(dfn['Category'][x])
            col.cells[9].text = str(dfn['CVE ID'][x])
            col.cells[10].text = str(dfn['Vendor Reference'][x])
            col.cells[11].text = str(dfn['Threat'][x])
            col.cells[12].text = str(dfn['Impact'][x])
            col.cells[13].text = str(dfn['Solution'][x])
            col.cells[14].text = str(dfn['Results'][x])
            col.cells[15].text = ''
        pp = pp + 1
        t = t + 1
    progress_bar.close()
    document.save(docx_filename)
    return docx_filename

def main():
    try:
        print_logo()

        filename = input(Fore.CYAN + "Enter CSV File Name (Ex:file.csv): " + Style.RESET_ALL)
        if not os.path.exists(filename):
                raise FileNotFoundError(f"{filename} file not found, please keep the CSV file in current directory")
    except Exception as e:
        print(Fore.RED + "[-]", str(e) + Style.RESET_ALL)
        exit()         
    print(Fore.GREEN + f"[+] CSV File: {filename} Loaded Successfully" + Style.RESET_ALL)
    print(Fore.GREEN + "[+] Converting CSV To XLSX" + Style.RESET_ALL)
    xlsx_file = merge_csv_to_xlsx(filename)
    print(Fore.GREEN + "[+] Conversion Completed. XLSX File Created: " + xlsx_file + Style.RESET_ALL)
    print(Fore.GREEN + "[+] Processing Data" + Style.RESET_ALL)
    df = process_data(xlsx_file)
    print(Fore.GREEN + "[+] Data Processing Completed" + Style.RESET_ALL)
    print(Fore.GREEN + "[+] Creating DOCX File" + Style.RESET_ALL)
    docx_file = create_docx_file(filename)
    print(Fore.GREEN + "[+] DOCX File Created: " + docx_file + Style.RESET_ALL)
    print(Fore.GREEN + "[+] Writing CSV Data To DOCX" + Style.RESET_ALL)
    write_csv_data_to_docx(df, docx_file)
    print(Fore.GREEN + "[+] Writing CSV Data To DOCX Completed" + Style.RESET_ALL)
    print(Fore.GREEN + "[+] Process Completed Successfully" + Style.RESET_ALL)   
       
if __name__ == "__main__":
    warnings.filterwarnings("ignore")
    try:
        main()
    except KeyboardInterrupt:
        print(Fore.RED + "\n[-] Program interrupted by the user. Exiting..." + Style.RESET_ALL)
        sys.exit(0)    
